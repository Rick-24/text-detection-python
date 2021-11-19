from docx import Document
import numpy as np
import pandas as pd

from text_detection.algorithm.jieba import vector_similarity, sentences_keyword_match


def document_analysis(filePath_input, filePath_rule, similarity_min, keyword_min):
    # 打开word文档
    document_input = Document(filePath_input)
    document_rule = Document(filePath_rule)

    policy_warehouse = []  # 以字典存储整个文章的每个整句话，字典属性：句子本身、关键字、第几章、第几条
    policy_userinput = []  # 以字典存储用户输入的文档中每句话的关键字，字典属性：句子本身、关键字、用户输入的第几条规则

    # 读取政策库所有政策
    all_paragraphs = document_rule.paragraphs  # 这是个列表
    duan = 0
    for paragraph in all_paragraphs:
        sentences = paragraph.text
        if sentences != "":  # 舍弃空句
            duan += 1
            sentences = sentences.split("。")
            for ju, sentence in enumerate(sentences):
                policy_warehouse.append({'sentence': sentence, 'duan': str(duan + 1), 'ju': str(ju + 1)})

    # 读取用户输入文档的查询语句
    all_paragraphs = document_input.paragraphs
    duan = 0
    for paragraph in all_paragraphs:
        sentences = paragraph.text
        if sentences != "":  # 舍弃空句
            duan += 1
            sentences = sentences.split("。")
            for ju, sentence in enumerate(sentences):
                policy_userinput.append({'sentence': sentence, 'duan': str(duan + 1), 'ju': str(ju + 1)})

    '''为匹配列表处理匹配度'''
    match_list = sentences_similarity(policy_warehouse, policy_userinput, similarity_min, keyword_min)  # 遍历两个文档以获取匹配序列

    def key_return(e):  # 定义一个返回'similarity'键的函数
        return e['similarity']

    match_list.sort(key=key_return, reverse=True)  # 倒序(由高到低)，且以字典的match_count的值来排序，排不排序应该都行，排完看着舒服
    # print(match_list)

    count_list = []  # 创建一个只存储match_list中match_count数据的一维数组
    for m in match_list:  # 我这里使用for循环创建的，有没有办法用一个函数就能取出来这一列？
        count_list.append(m['similarity'])
    count_list = np.array(count_list)  # list转array
    try:
        count_list = pd.cut(count_list, bins=3, labels=False)  # 数据自动分箱，分箱(bins=3个箱子)操作后，返回值为0,1,2的数组，表示"低,中,高"
    except:
        pass
    if (len(count_list) != 0):
        print("输入文档与规则库文件" + filePath_rule + "的分箱结果:" + str(count_list))
        # 看一下分箱结果吧！# pd.cut函数设置参数retbins=True，函数就可以返回分箱的边界值
    else:
        print("输入文档未在此规则库文件中找到结果:" + filePath_rule)
    for index, m in enumerate(match_list):
        m['matching_degree'] = str(count_list[index])  # 为match_list增加"匹配度"键值对，2为高，1为中，0为低

    match_list = {'filepath_input': filePath_input, 'filepath_rule': filePath_rule, 'match_list': match_list}
    return match_list


# 遍历政策库和用户输入所有句子（遍历两个文档以获取匹配序列）
def sentences_similarity(policy_warehouse, policy_userinput, similarity_min, keyword_min):
    match_list = []
    for s1 in policy_warehouse:
        for s2 in policy_userinput:
            try:
                if s1['sentence'] == "" or s2['sentence'] == "":  # 句子为空
                    continue
                else:
                    similarity = vector_similarity(s1['sentence'], s2['sentence'])
                    if similarity >= similarity_min:  # 设置相似度门槛，高于门槛才加入list
                        keyword_match = sentences_keyword_match(s1['sentence'], s2['sentence'])
                        if keyword_match >= keyword_min:  # 设置关键字匹配数门槛，高于门槛才加入list
                            match_list.append({'warehouse': s1, 'userinput': s2, 'similarity': similarity,
                                               'keyword_match': keyword_match})
            except (KeyError, TypeError):
                print("Error")
    return match_list
