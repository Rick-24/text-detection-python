# 加载模型
import gensim
import jieba
import numpy as np
from scipy.linalg import norm
import time

time_start = time.time()

model = gensim.models.Word2Vec.load('D:/全部资料/研究生学前学习/政策NLP demo/training data/zh/wiki.zh.text.model')
print(model.get_latest_training_loss())

time_end = time.time()

print('Initializing time cost = %fs' % (time_end - time_start))
print(model)

#使用word2vec对两个句子的相似度进行分析
def vector_similarity(s1, s2):
    def sentence_vector(s):
        words = jieba.lcut(s)
        v = np.zeros(model.vector_size)  # 注意要与model的vector的size保持一直，否则无法进行array的计算
        for word in words:
            try:
                v += model.wv[word]
            except:
                pass
        v /= len(words)
        return v
    
    v1, v2 = sentence_vector(s1), sentence_vector(s2)
    return np.dot(v1, v2) / (norm(v1) * norm(v2))

#使用jieba关键字对两个句子的相似度进行分析
def sentences_keyword_match(s1, s2): 
    tr_result1 = analyse.textrank(s1, topK=10) # 分析关键字
    tr_result2 = analyse.textrank(s2, topK=10) # 分析关键字
    count = 0
    for keyword1 in tr_result1:
        for keyword2 in tr_result2:
            if keyword1 == keyword2:
                count += 1
    return count

from urllib.parse import unquote, quote
from docx import Document
import jieba
import jieba.posseg as posseg
import jieba.analyse as analyse
import re
import numpy as np
import pandas as pd

import json
import urllib.request
import threading
import socket

def document_analysis(filePath_input, filePath_rule, similarity_min, keyword_min):
    #打开word文档
    document_input = Document(filePath_input)
    document_rule = Document(filePath_rule)
    
    policy_warehouse = [] # 以字典存储整个文章的每个整句话，字典属性：句子本身、关键字、第几章、第几条
    policy_userinput = [] # 以字典存储用户输入的文档中每句话的关键字，字典属性：句子本身、关键字、用户输入的第几条规则
    
    #获取所有段落
    all_paragraphs = document_rule.paragraphs # 这是个列表
    #读取政策库所有政策
    for duan, paragraph in enumerate(all_paragraphs):
        sentences = paragraph.text
        sentences = sentences.split("。")
        for ju, sentence in enumerate(sentences):
            if sentence != "":
                policy_warehouse.append({'sentence': sentence, 'duan': str(duan+1), 'ju': str(ju+1)})

    all_paragraphs = document_input.paragraphs
    
    #读取用户输入文档的查询语句
    for duan, paragraph in enumerate(all_paragraphs):
        sentences = paragraph.text
        sentences = sentences.split("。")
        for ju, sentence in enumerate(sentences):
            if sentence != "": # 舍弃空句
                policy_userinput.append({'sentence': sentence, 'duan': str(duan+1), 'ju': str(ju+1)})
 
    '''为匹配列表处理匹配度'''
    match_list = sentences_similarity(policy_warehouse, policy_userinput, similarity_min, keyword_min) # 遍历两个文档以获取匹配序列
    def key_return(e): # 定义一个返回'similarity'键的函数
        return e['similarity']
    match_list.sort(key=key_return, reverse=True) # 倒序(由高到低)，且以字典的match_count的值来排序，排不排序应该都行，排完看着舒服
    #print(match_list)
    
    count_list = [] # 创建一个只存储match_list中match_count数据的一维数组
    for m in match_list: # 我这里使用for循环创建的，有没有办法用一个函数就能取出来这一列？
        count_list.append(m['similarity'])
    count_list = np.array(count_list) # list转array
    try:
        count_list = pd.cut(count_list, bins=3, labels=False) # 数据自动分箱，分箱(bins=3个箱子)操作后，返回值为0,1,2的数组，表示"低,中,高"
    except:
        pass
    if(len(count_list) != 0):
        print("输入文档与规则库文件"+filePath_rule+"的分箱结果:"+str(count_list)) 
        # 看一下分箱结果吧！# pd.cut函数设置参数retbins=True，函数就可以返回分箱的边界值
    else:
        print("输入文档未在此规则库文件中找到结果:" + filePath_rule)
    for index, m in enumerate(match_list):
        m['matching_degree'] = str(count_list[index]) # 为match_list增加"匹配度"键值对，2为高，1为中，0为低
        
    match_list = {'filepath_input': filePath_input, 'filepath_rule': filePath_rule, 'match_list': match_list}
    return match_list

# 遍历政策库和用户输入所有句子（遍历两个文档以获取匹配序列）
def sentences_similarity(policy_warehouse, policy_userinput, similarity_min, keyword_min):
    match_list = []
    for s1 in policy_warehouse:
        for s2 in policy_userinput:
            try:
                if (s1['sentence'] == "" or s2['sentence'] == ""): # 句子为空
                    continue
                else:
                    similarity = vector_similarity(s1['sentence'], s2['sentence'])
                    if similarity >= similarity_min: #设置相似度门槛，高于门槛才加入list
                        keyword_match = sentences_keyword_match(s1['sentence'], s2['sentence'])
                        if keyword_match >= keyword_min: #设置关键字匹配数门槛，高于门槛才加入list
                            match_list.append({'warehouse': s1, 'userinput': s2, 'similarity': similarity, 'keyword_match': keyword_match})
            except (KeyError, TypeError):
                print("Error")
    return match_list

class Request:
    def __init__(self, r):
        self.content = r
        self.method = r.split()[0]
        self.path = r.split()[1]
        self.body = r.split('\r\n\r\n', 1)[1]
 
    def form_body(self):
        return self._parse_parameter(self.body)
 
    def parse_path(self):
        index = self.path.find('?')
        if index == -1:
            return self.path, {}
        else:
            path, query_string = self.path.split('?', 1)
            query = self._parse_parameter(query_string)
            return path, query
 
    @property
    def headers(self):
        header_content = self.content.split('\r\n\r\n', 1)[0].split('\r\n')[1:]
        result = {}
        for line in header_content:
            k, v = line.split(': ')
            result[quote(k)] = quote(v)
        return result
 
    @staticmethod
    def _parse_parameter(parameters):
        args = parameters.split('&')
        query = {}
        for arg in args:
            k, v = arg.split('=')
            query[k] = unquote(v)
        return query

def multithreading_simulate(): # 测试函数，不停调用document_analysis()分析函数，没用多线程
    import os
    
    input_folder = "D:/全部资料/研究生学前学习/政策NLP demo/rule/"
    fileName_input = "/input/检测文件-甘井子区经济工作奖励暂行办法.docx"
    filePath_input = input_folder + fileName_input
    
    rule_folder = 'D:/全部资料/研究生学前学习/政策NLP demo/rule/21/02/'
    fileName_rule = []
    for i in os.walk(rule_folder):
        if len(i[1]) != 0:   # 不包含子文件夹
            fileName_rule.extend(i[2])
    
    print(fileName_rule)
    match_list_total = []
    
    for f_name in fileName_rule:
        filePath_rule = rule_folder + f_name
        match_list = document_analysis(filePath_input, filePath_rule, 0.9, 3) ######调用两个文档的分析函数######
        match_list_total.append(match_list)

    return match_list_total
    
'''''
建立一个python server，监听指定端口，
如果该端口被远程连接访问，则获取远程连接，然后接收数据，
并且做出相应反馈。
'''
if __name__=="__main__":
#     import socket
#     print("Server is starting")
#     sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
#     sock.bind(('localhost', 2452)) #配置soket，绑定IP地址和端口号
#     sock.listen(5) #设置最大允许连接数，各连接和server的通信遵循FIFO原则
#     print("Server is listenting port 2452, with max connection 5")
#     while True: #循环轮询socket状态，等待访问
#         connection,address = sock.accept()
#         try:
#             connection.settimeout(50)
#             #获得一个连接，然后开始循环处理这个连接发送的信息
#             '''''
#             如果server要同时处理多个连接，则下面的语句块应该用多线程来处理，
#             否则server就始终在下面这个while语句块里被第一个连接所占用，
#             无法去扫描其他新连接了，但多线程会影响代码结构，所以记得在连接数大于1时
#             下面的语句要改为多线程即可。
#             '''
            
#             buf = connection.recv(1024)
#             input = bytes.decode(buf, "utf-8")
# #             req = Request(input)
# #             print("body:"+req.body)
            
#             print(input)
#             requestBody = json.loads(input)
#             filePath = requestBody['filePath']
                
#             if len(filePath) != 0:
#                 match_list = document_analysis(filePath)
#                 print("\n\n\n matchlist:")
#                 print(match_list)
#                 resp =json.dumps(match_list)
#                 connection.send(resp.encode(encoding='UTF-8'))
#             else:
#                 print("Bad Request !!!")
#                 break #退出连接监听循环
#         except socket.timeout: #如果建立连接后，该连接在设定的时间内无数据发来，则time out
#              print('time out')
#         print("closing one connection") #当一个连接监听循环退出后，连接可以关掉
#         connection.close()
    
    '''测试单个规则文件'''
    filePath_input = "D:/全部资料/研究生学前学习/政策NLP demo/rule//input/检测文件-甘井子区经济工作奖励暂行办法.docx"
    filePath_rule = "D:/全部资料/研究生学前学习/政策NLP demo/rule/21/02/中共大连市委大连市人民政府关于促进物流业发展的实施意见.docx"
    similarity_min = 0.9 # 设置两句子相似度最低值
    keyword_min = 3      # 设置两句子关键字匹配数最低值
    match_list = document_analysis(filePath_input, filePath_rule, similarity_min, keyword_min)

    '''测试一个目录下的一组规则文件'''
    # match_list = multithreading_simulate()
    # print("\n\n输入文档与这个目录下的所有文件的匹配结果：")

    print(match_list)